import os
from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
from collections import Counter



def remove_repeated_lines(documents, min_repetitions=2):
    """
    Removes repeated lines that appear across multiple pages.
    Useful for removing headers, footers, ads, cookie notices, and navigation text.
    """
    line_counter = Counter()

    # Count how often each line appears across all pages
    for doc in documents:
        lines = [
            line.strip()
            for line in doc.page_content.splitlines()
            if line.strip()
        ]

        unique_lines_on_page = set(lines)

        for line in unique_lines_on_page:
            if len(line) > 5:
                line_counter[line] += 1

    repeated_lines = {
        line
        for line, count in line_counter.items()
        if count >= min_repetitions
    }

    cleaned_documents = []

    for doc in documents:
        lines = [
            line.strip()
            for line in doc.page_content.splitlines()
            if line.strip()
        ]

        cleaned_lines = [
            line
            for line in lines
            if line not in repeated_lines
        ]

        doc.page_content = "\n".join(cleaned_lines)

        cleaned_text = "\n".join(cleaned_lines)

        ad_phrases = [
            "Advertisement",
            "ADVERTISEMENT",
            "Sponsored Content",
            "Subscribe",
        ]

        for phrase in ad_phrases:
            cleaned_text = cleaned_text.replace(phrase, "")

        doc.page_content = cleaned_text
        cleaned_documents.append(doc)

    print(f"Removed repeated lines: {len(repeated_lines)}")

    return cleaned_documents


def load_pdf(file_path: str):
    loader = PyPDFLoader(file_path)
    documents = loader.load()

    documents = remove_repeated_lines(documents, min_repetitions=2)

    for page_number, doc in enumerate(documents, start=1):
        doc.metadata["page"] = page_number
        doc.metadata["source"] = file_path

        print(f"Page {page_number} text length: {len(doc.page_content.strip())}")

    return documents


def load_txt(file_path: str):
    with open(file_path, "r", encoding="utf-8") as file:
        text = file.read()

    document = Document(
        page_content=text,
        metadata={
            "source": file_path,
            "page": 1
        }
    )

    print(f"TXT text length: {len(text.strip())}")

    return [document]


def load_docx(file_path: str):
    docx_file = DocxDocument(file_path)

    paragraphs = [
        paragraph.text
        for paragraph in docx_file.paragraphs
        if paragraph.text.strip()
    ]

    text = "\n".join(paragraphs)

    document = Document(
        page_content=text,
        metadata={
            "source": file_path,
            "page": 1
        }
    )

    print(f"DOCX text length: {len(text.strip())}")

    return [document]


def load_document(file_path: str):
    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return load_pdf(file_path)

    if extension == ".txt":
        return load_txt(file_path)

    if extension == ".docx":
        return load_docx(file_path)

    raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")